from __future__ import annotations

import csv
import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "LLM辅助Linux内核缺陷修复_结项论文.docx"
ASSET_DIR = ROOT / "docs" / "_paper_assets"

BLUE = "24507A"
DARK = "17324D"
LIGHT = "EAF1F7"
PALE = "F4F7FA"
GRAY = "6B7280"
INK = "1F2937"
WHITE = "FFFFFF"
GOLD = "C08A24"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, east_asia="宋体", ascii_font="Calibri", size=11, bold=None, color=None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_keep(paragraph, keep_next=False, keep_lines=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=9, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "目录将在打开文档时自动更新。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=DARK)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_para_keep(p, keep_next=True)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", size=9.5, color=GRAY)
    set_para_keep(p, keep_next=True)
    return p


def add_picture(doc: Document, path: Path, caption: str, width=6.05):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    set_para_keep(p, keep_next=True)
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], caption: str):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for index, text in enumerate(headers):
        cell = hdr.cells[index]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, east_asia="微软雅黑", size=9.2, bold=True, color=WHITE)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 or len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(value)
            set_run_font(run, size=9.2)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def font(size: int, bold=False):
    paths = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in paths:
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size, index=0)
            except Exception:
                continue
    return ImageFont.load_default()


def draw_flow(path: Path) -> None:
    image = Image.new("RGB", (1800, 930), "white")
    draw = ImageDraw.Draw(image)
    title = font(48, True)
    label = font(34, True)
    body = font(28)
    small = font(24)
    draw.text((900, 50), "LLM辅助Linux内核缺陷修复实验流程", font=title, fill="#17324D", anchor="ma")
    xs = [100, 430, 760, 1090, 1420]
    labels = [
        ("崩溃输入", "Crash report\n源码片段\n可选 trace"),
        ("根因假设", "生成候选假设\n关联栈与源码证据"),
        ("自反思", "一致性检查\n选择最优假设"),
        ("补丁生成", "Baseline / Trace\nSemantic Guard"),
        ("分层验证", "静态检查\n源码适用\n编译与动态验证"),
    ]
    for i, (head, detail) in enumerate(labels):
        x = xs[i]
        draw.rounded_rectangle((x, 250, x + 270, 620), radius=28, fill="#EAF1F7", outline="#24507A", width=4)
        draw.rounded_rectangle((x, 250, x + 270, 330), radius=28, fill="#24507A")
        draw.rectangle((x, 300, x + 270, 330), fill="#24507A")
        draw.text((x + 135, 292), head, font=label, fill="white", anchor="mm")
        lines = detail.split("\n")
        for j, line in enumerate(lines):
            draw.text((x + 135, 405 + j * 62), line, font=body, fill="#1F2937", anchor="mm")
        if i < len(labels) - 1:
            draw.line((x + 278, 435, xs[i + 1] - 12, 435), fill="#C08A24", width=8)
            draw.polygon([(xs[i + 1] - 12, 435), (xs[i + 1] - 38, 420), (xs[i + 1] - 38, 450)], fill="#C08A24")
    draw.rounded_rectangle((210, 720, 1590, 840), radius=24, fill="#F4F7FA", outline="#A9BACB", width=3)
    draw.text((900, 760), "核心改进：在补丁阶段加入 Semantic Guard，显式约束模型不得删除、注释或绕过核心功能", font=small, fill="#17324D", anchor="ma")
    draw.text((900, 805), "证据边界：人工标签 ≠ 真实修复；仅将已执行的 apply、compile 和 reproducer 结果计入工程验证", font=small, fill="#6B7280", anchor="ma")
    image.save(path, quality=95)


def draw_group_chart(path: Path) -> None:
    image = Image.new("RGB", (1700, 1000), "white")
    draw = ImageDraw.Draw(image)
    title = font(46, True)
    axis = font(25)
    label = font(28, True)
    draw.text((850, 45), "三组实验人工评价结果", font=title, fill="#17324D", anchor="ma")
    groups = ["Baseline", "With Trace", "Improved"]
    data = {"plausible": [1, 0, 1], "helpful": [1, 3, 3], "incorrect": [6, 5, 4]}
    colors = {"plausible": "#2F6B4F", "helpful": "#C08A24", "incorrect": "#A24B4B"}
    x0, y0, chart_w, chart_h = 170, 150, 1320, 650
    draw.line((x0, y0, x0, y0 + chart_h), fill="#44515F", width=3)
    draw.line((x0, y0 + chart_h, x0 + chart_w, y0 + chart_h), fill="#44515F", width=3)
    for tick in range(0, 9, 2):
        y = y0 + chart_h - tick / 8 * chart_h
        draw.line((x0 - 10, y, x0 + chart_w, y), fill="#D8E0E8", width=2)
        draw.text((x0 - 25, y), str(tick), font=axis, fill="#4B5563", anchor="rm")
    group_span = chart_w / 3
    bar_w = 80
    for gi, group in enumerate(groups):
        center = x0 + group_span * (gi + 0.5)
        for bi, key in enumerate(["plausible", "helpful", "incorrect"]):
            value = data[key][gi]
            height = value / 8 * chart_h
            x = center + (bi - 1) * 100 - bar_w / 2
            draw.rounded_rectangle((x, y0 + chart_h - height, x + bar_w, y0 + chart_h), radius=10, fill=colors[key])
            draw.text((x + bar_w / 2, y0 + chart_h - height - 14), str(value), font=label, fill=colors[key], anchor="ms")
        draw.text((center, y0 + chart_h + 55), group, font=label, fill="#1F2937", anchor="ma")
    legend_x = 420
    for i, key in enumerate(["plausible", "helpful", "incorrect"]):
        x = legend_x + i * 330
        draw.rounded_rectangle((x, 900, x + 34, 934), radius=6, fill=colors[key])
        draw.text((x + 50, 917), key, font=axis, fill="#374151", anchor="lm")
    image.save(path, quality=95)


def draw_evidence_chart(path: Path) -> None:
    image = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(image)
    title = font(46, True)
    head = font(31, True)
    body = font(25)
    draw.text((900, 45), "分层验证证据链", font=title, fill="#17324D", anchor="ma")
    layers = [
        ("人工语义评价", "24/24", "plausible 2，helpful 7，incorrect 15", "#EAF1F7", "#24507A"),
        ("真实源码适用", "24/24", "模型补丁 apply 通过 10/24", "#EEF5EF", "#2F6B4F"),
        ("局部编译", "1/24", "bug_008 improved：fs/namespace.o 通过", "#FFF5E6", "#C08A24"),
        ("动态复现", "1/24", "parent 约 59 秒复现；patched 3×6 分钟 clean pass", "#F8ECEC", "#A24B4B"),
    ]
    for i, (name, coverage, detail, fill, outline) in enumerate(layers):
        top = 150 + i * 205
        left = 180 + i * 100
        right = 1620 - i * 100
        draw.rounded_rectangle((left, top, right, top + 145), radius=26, fill=fill, outline=outline, width=4)
        draw.text((left + 45, top + 48), name, font=head, fill=outline, anchor="lm")
        draw.text((right - 45, top + 48), coverage, font=head, fill=outline, anchor="rm")
        draw.text(((left + right) / 2, top + 102), detail, font=body, fill="#374151", anchor="mm")
        if i < len(layers) - 1:
            draw.line((900, top + 150, 900, top + 198), fill="#7A8A99", width=6)
            draw.polygon([(900, top + 198), (884, top + 175), (916, top + 175)], fill="#7A8A99")
    draw.text((900, 1010), "越向下证据越强，但当前覆盖越小；未执行项记为 not_run，而不是失败。", font=body, fill="#6B7280", anchor="ms")
    image.save(path, quality=95)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.first_line_indent = Pt(22)

    heading_specs = {
        1: (16, 18, 10, BLUE),
        2: (13, 12, 6, BLUE),
        3: (12, 8, 4, DARK),
    }
    for level, (size, before, after, color) in heading_specs.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.first_line_indent = Pt(0)

    if "Abstract Text" not in [style.name for style in styles]:
        abstract = styles.add_style("Abstract Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        abstract = styles["Abstract Text"]
    abstract.font.name = "Calibri"
    abstract._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    abstract.font.size = Pt(10.5)
    abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.space_after = Pt(7)
    abstract.paragraph_format.line_spacing = 1.25
    abstract.paragraph_format.first_line_indent = Pt(21)


def setup_document() -> Document:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = "LLM辅助Linux内核缺陷修复 · 结项论文"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run, east_asia="微软雅黑", size=8.5, color=GRAY)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D6DEE6")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_number(section.footer.paragraphs[0])

    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    return doc


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("结 项 论 文")
    set_run_font(run, east_asia="黑体", size=20, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("LLM辅助Linux内核缺陷修复")
    set_run_font(run, east_asia="黑体", size=25, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    run = p.add_run("基于根因假设、自反思与语义保护约束的轻量化实验研究")
    set_run_font(run, east_asia="微软雅黑", size=13, color=GRAY)

    labels = ["学生姓名", "学号", "专业/班级", "指导教师", "完成日期"]
    values = ["________________", "________________", "________________", "________________", "2026年6月"]
    for label, value in zip(labels, values):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        r1 = p.add_run(f"{label}：  ")
        set_run_font(r1, east_asia="黑体", size=11, bold=True, color=DARK)
        r2 = p.add_run(value)
        set_run_font(r2, east_asia="宋体", size=11)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("软件质量保证课程项目")
    set_run_font(run, east_asia="微软雅黑", size=10, color=GRAY)
    doc.add_page_break()


def add_abstracts(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("摘  要")
    set_run_font(run, east_asia="黑体", size=16, bold=True, color=DARK)
    p.paragraph_format.space_after = Pt(14)

    abstracts = [
        "Linux内核包含大规模底层C代码、并发控制与硬件交互。Syzkaller能够持续发现真实内核崩溃，但其报告通常由调用栈、寄存器和sanitizer日志构成，缺少自然语言缺陷描述，导致根因定位与补丁构造高度依赖专家经验。近年来，大语言模型在代码理解和自动程序修复方面表现出潜力，但现有研究多集中于用户态仓库，Linux内核场景仍面临上下文规模大、错误语义隐含、验证成本高以及模型可能通过删除功能规避崩溃等问题。",
        "本文设计并实现了一套面向Linux内核缺陷的轻量化LLM修复流程。该流程以CrashFixer的开发者式调试思想为基础，将修复划分为根因假设生成、自反思和补丁生成三个阶段，并在补丁阶段加入Semantic Guard，明确禁止删除、注释或绕过核心功能，要求模型解释每个修改对原有语义的保持方式。实验选取8个真实kBenchSyz样本，覆盖内存泄漏、越界访问和空指针三类缺陷，设置Baseline、With Trace和Improved三组，共得到24个模型输出。",
        "评价采用分层证据体系。人工语义评价结果为plausible 2/24、helpful 7/24、incorrect 15/24；plausible与helpful合计比例由Baseline的25.0%提升至With Trace的37.5%，Improved达到50.0%。在真实Linux parent commit上，模型补丁应用通过率分别为37.5%、25.0%和62.5%。代表案例bug_008中，Improved补丁与开发者补丁内容一致，三路局部编译均通过；原始版本约59秒复现目标KASAN崩溃，而修复版本连续3次、每次6分钟动态运行均未出现目标崩溃、特殊崩溃或worker异常。",
        "结果表明，执行轨迹能够提升根因方向的有用性，Semantic Guard有助于减少无编辑输出并促使模型生成较小、可应用的修改，但不能替代完整的资源所有权分析。两个内存泄漏样本的6个输出全部错误，说明局部源码和短轨迹不足以稳定恢复跨路径cleanup语义。本文还借鉴RGym的轻量验证思想，扩展记录定位命中、补丁适用、编译、崩溃消除、开发者补丁相似度、API成本和重试次数，为后续更大规模实验提供可复现基础。",
    ]
    for text in abstracts:
        p = doc.add_paragraph(style="Abstract Text")
        run = p.add_run(text)
        set_run_font(run, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("关键词：")
    set_run_font(run, east_asia="黑体", size=10.5, bold=True, color=DARK)
    run = p.add_run("大语言模型；Linux内核；自动程序修复；Syzkaller；Semantic Guard；动态验证")
    set_run_font(run, size=10.5)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ABSTRACT")
    set_run_font(run, size=16, bold=True, color=DARK)
    p.paragraph_format.space_after = Pt(14)
    english = [
        "Linux kernel crash reports generated by fuzzers contain stack traces and sanitizer logs but usually lack natural-language bug descriptions. Repairing them requires system-level reasoning, broad code context, and expensive compile-and-run validation. This project develops a lightweight LLM-assisted workflow for real Linux kernel defects.",
        "The workflow contains hypothesis generation, self-reflection, and patch generation. A Semantic Guard is added to the final stage to prohibit deletion, commenting-out, or bypassing of core functionality and to require an explicit semantic-preservation explanation. Eight real kBenchSyz cases are evaluated under Baseline, With Trace, and Improved settings, producing 24 outputs.",
        "Manual analysis yields 2 plausible, 7 helpful, and 15 incorrect outputs. The useful-output ratio increases from 25.0% in Baseline to 37.5% with trace information and 50.0% with Semantic Guard. Real-source patch applicability reaches 62.5% in the Improved group. For bug_008, the generated patch matches the developer fix, passes local compilation, and completes three independent six-minute dynamic validation runs without the target crash. The results support the value of structured reasoning and semantic constraints, while also revealing clear limitations on memory-leak and ownership-related defects.",
    ]
    for text in english:
        p = doc.add_paragraph(style="Abstract Text")
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        set_run_font(run, east_asia="宋体", ascii_font="Calibri", size=10.5)
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    set_run_font(run, bold=True, size=10.5, color=DARK)
    run = p.add_run("large language model; Linux kernel; automated program repair; Syzkaller; semantic guard; dynamic validation")
    set_run_font(run, size=10.5)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目  录")
    set_run_font(run, east_asia="黑体", size=16, bold=True, color=DARK)
    p.paragraph_format.space_after = Pt(14)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Pt(0)
    add_toc(toc)
    doc.add_page_break()


def build_content(doc: Document, figures: dict[str, Path]) -> None:
    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_body(doc, "Linux内核是服务器、移动设备、嵌入式系统和云平台的重要基础软件。其代码规模超过两千万行，涉及内存管理、文件系统、设备驱动、网络协议栈和并发同步等多个子系统。内核缺陷一旦触发，可能造成系统崩溃、数据损坏、权限绕过或拒绝服务，因此缺陷修复不仅是功能维护问题，也是软件质量与安全保障问题。")
    add_body(doc, "Syzkaller通过覆盖引导的系统调用序列生成持续发现内核异常，并由syzbot自动完成构建、复现和报告。与普通GitHub issue不同，内核fuzzer报告主要包含KASAN、UBSAN等sanitizer信息、调用栈和复现程序，缺少对业务意图和预期行为的自然语言解释。开发者需要从低层日志中推断失效路径、对象生命周期和锁语义，修复门槛较高。")
    add_body(doc, "大语言模型具备代码生成和跨文本推理能力，为自动程序修复提供了新的实现路径。然而，直接将通用代码模型用于内核存在三类风险：其一，模型可能仅围绕崩溃行生成局部补丁，忽略跨函数状态；其二，补丁可能通过提前返回或删除功能让复现器不再触发，却破坏原有语义；其三，未经编译和动态验证的补丁容易被误判为成功。因此，本项目关注的不只是“能否生成补丁”，还关注“补丁是否在证据边界内可信”。")

    add_heading(doc, "1.2 研究目标与问题", 2)
    add_body(doc, "本项目目标是构建一套本科课程项目规模内可执行、可复现的Linux内核LLM修复实验流程，并分析执行轨迹与语义保护约束的作用。围绕该目标，本文提出以下研究问题：")
    add_body(doc, "RQ1：在只提供崩溃报告和局部源码时，LLM能够生成多少具有根因参考价值或补丁价值的输出？", "RQ1：")
    add_body(doc, "RQ2：加入人工整理的trace summary后，根因判断与补丁质量是否改善？", "RQ2：")
    add_body(doc, "RQ3：Semantic Guard能否减少无编辑或功能破坏型修复，并提高真实源码上的补丁适用率？", "RQ3：")
    add_body(doc, "RQ4：人工语义评价、源码适用、编译和动态验证之间存在多大差异，如何避免将轻量指标误称为真实修复率？", "RQ4：")

    add_heading(doc, "1.3 主要工作与贡献", 2)
    add_body(doc, "本文完成了以下工作。第一，基于CrashFixer思想实现根因假设、自反思和补丁生成三阶段流程，并支持DeepSeek接口和结构化JSON输出。第二，从kBenchSyz中固定筛选8个真实内核缺陷，建立可复现实验数据目录。第三，在补丁阶段加入Semantic Guard，对删除核心功能、无依据提前返回和语义绕过进行显式约束。第四，建立从人工评价到真实源码apply、局部编译和QEMU动态复现的分层证据链。第五，借鉴RGym的评测思路新增定位命中、开发者补丁相似度、API成本与重试次数字段，使后续实验能够同时报告质量、工程可行性和资源消耗。")

    add_heading(doc, "1.4 论文结构", 2)
    add_body(doc, "第2章介绍自动程序修复、LLM修复代理以及Linux内核相关基准；第3章描述本文方法和Semantic Guard；第4章给出系统实现与评价体系；第5章报告实验结果；第6章分析代表案例；第7章讨论有效性、局限性与RGym带来的改进方向；第8章总结全文。")

    add_heading(doc, "2 相关工作", 1)
    add_heading(doc, "2.1 自动程序修复", 2)
    add_body(doc, "传统自动程序修复通常将修复建模为搜索或合成问题。GenProg使用遗传编程在候选修改空间中搜索可通过测试的程序变体[1]；随后，基于约束求解和神经机器翻译的方法逐步提高补丁生成能力。CoCoNuT通过上下文感知的神经翻译模型进行程序修复[2]。这些方法推动了APR发展，但常依赖规模较小、测试充分且错误位置较明确的基准。")
    add_body(doc, "代码大模型将预训练代码知识与自然语言推理结合，使修复系统可以阅读issue、调用工具并跨文件编辑。ChatGPT在QuixBugs上的研究显示了对话式修复的潜力[3]，SWE-bench则将评测推进到真实GitHub仓库级问题[4]。不过，用户态仓库的构建和测试模式与Linux内核仍有明显差异。")

    add_heading(doc, "2.2 Linux内核修复基准", 2)
    add_body(doc, "kGym及其kBench数据集首次系统提供面向Linux内核崩溃修复的实验平台和真实数据。样本包含崩溃栈、复现程序、开发者修复和相关元数据，平台支持并行编译内核、运行虚拟机、检测崩溃和应用补丁[5]。其基线结果表明，即使向模型披露缺陷文件，修复率仍然很低，说明内核语义、并发和硬件相关代码对LLM构成显著挑战。")
    add_body(doc, "CrashFixer进一步模拟内核开发者工作流，先显式生成根因假设，再生成补丁，并扩展kGym为kGymSuite以支持大规模代理实验[6]。论文显示，复杂系统中的假设生成有助于修复，但人工评价中仍存在较多incorrect补丁，其中包括通过删除或绕过功能消除崩溃的修复。本文的Semantic Guard直接针对这一风险。")

    add_heading(doc, "2.3 RGym及新近研究", 2)
    add_body(doc, "RGym提出可在普通本地硬件上运行的轻量、平台无关内核APR评测框架，并使用调用栈和blamed commit进行更现实的定位。其消融实验同时比较定位策略、Prompt、模型和反馈重试，并报告通过率与单缺陷成本[7]。RGym与本项目都强调降低验证门槛，但本文的重点不同：RGym主要改进定位和评测基础设施，本文重点分析trace与Semantic Guard对补丁行为的影响。")
    add_body(doc, "Code Researcher通过代码库语义、模式和提交历史进行多步检索，说明大型系统修复依赖全局上下文[8]；SemAgent强调问题语义、代码语义与执行语义的一致性，指出超局部修复容易过拟合[9]；Live-kBench进一步揭示静态基准的数据污染风险，并用持续更新的内核缺陷评估代理[10]。这些工作共同说明：内核APR的关键不只是更强模型，还包括定位、上下文、反馈和可信验证。")

    add_heading(doc, "2.4 研究空白与本文定位", 2)
    add_body(doc, "现有工作已经证明LLM可以参与Linux内核修复，因此本文不主张“首次使用LLM修复内核”。本文定位为：在有限计算资源和小样本条件下，复现CrashFixer的核心推理流程，加入可审计的语义保护约束，并使用RGym式轻量指标和分层工程证据评估结果。该定位更符合课程项目规模，也能形成明确、可验证的创新点。")

    add_heading(doc, "3 方法设计", 1)
    add_heading(doc, "3.1 总体流程", 2)
    add_body(doc, "实验输入由crash report、localization file对应的源码片段以及可选trace summary组成。系统依次执行Hypothesis Generation、Self-Reflection和Patch Generation。每个阶段要求模型返回严格JSON，原始响应与解析结果分别保存，便于复核。")
    add_picture(doc, figures["flow"], "图3-1 本项目三阶段修复流程与分层验证", width=6.15)

    add_heading(doc, "3.2 根因假设生成", 2)
    add_body(doc, "第一阶段要求模型结合崩溃类型、调用栈和源码提出多个候选根因。每个假设包含可疑函数或变量、根因解释、证据、修复思路和不确定性。显式假设的作用是把“直接改代码”转换为“先解释失效机制再修改”，从而减少仅依据最后一层栈帧进行机械防御的倾向。")

    add_heading(doc, "3.3 自反思", 2)
    add_body(doc, "第二阶段逐一检查候选假设与崩溃报告、栈轨迹和源码是否一致，记录优势、弱点和证据一致性，并选择一个最可信假设。该阶段并不保证推理正确，但为错误假设提供了第二次筛选机会，并把最终补丁与选定根因绑定。")

    add_heading(doc, "3.4 补丁生成与Semantic Guard", 2)
    add_body(doc, "第三阶段采用replace-based edit格式，每个edit明确给出文件、原始片段、替换片段和理由。Improved组额外加入Semantic Guard，其约束包括：不得删除或注释核心功能；不得通过无依据提前返回绕过正常路径；对每个修改说明为何保持原有功能；优先采用最小局部修改；若证据不足，应明确限制而不是伪造上下文。")
    add_body(doc, "Semantic Guard不是形式化语义证明。它是一种Prompt层防护，用于提高补丁意图的可审计性，并减少明显的截肢式修复。真正的正确性仍需通过源码、编译和动态证据逐层建立。")

    add_heading(doc, "3.5 实验分组", 2)
    add_table(
        doc,
        ["组别", "输入", "补丁阶段", "目的"],
        [
            ["Baseline", "crash report + source", "普通Patch Prompt", "建立基础性能"],
            ["With Trace", "Baseline + trace summary", "普通Patch Prompt", "观察轨迹信息作用"],
            ["Improved", "Baseline + trace summary", "Semantic Guard Prompt", "观察语义约束作用"],
        ],
        [1600, 3000, 2300, 2460],
        "表3-1 三组实验设计",
    )

    add_heading(doc, "4 系统实现与评价体系", 1)
    add_heading(doc, "4.1 项目结构", 2)
    add_body(doc, "项目由data、prompts、src、outputs、results和docs六部分组成。data/selected保存固定样本；prompts保存四类模板；src包含加载、Prompt构造、模型调用、静态检查和统计脚本；outputs按实验组保存每个阶段的输入输出；results保存人工评价、源码核验、编译和动态验证证据；docs保存复现指南和分析材料。该结构使数据、方法、原始响应和评价结果能够相互追溯。")

    add_heading(doc, "4.2 数据集与样本选择", 2)
    add_body(doc, "实验从kBenchSyz中筛选8个状态已修复、developer patch较短且主要涉及单文件的样本。样本覆盖2个内存泄漏、3个越界访问和3个空指针问题。source.c由developer patch context重构，仅用于轻量Prompt实验；真实适用性核验另行checkout对应parent commit，避免把片段匹配等同于真实源码应用。")
    add_table(
        doc,
        ["编号", "类型", "子系统", "定位文件"],
        [
            ["bug_001", "memory leak", "netfilter", "net/netfilter/nf_tables_api.c"],
            ["bug_002", "memory leak", "media/usb", "drivers/media/usb/dvb-usb/cinergyT2-core.c"],
            ["bug_003", "out-of-bounds", "printk", "kernel/printk/printk.c"],
            ["bug_004", "out-of-bounds", "xfrm/selinux", "net/xfrm/xfrm_user.c"],
            ["bug_005", "out-of-bounds", "usbip", "drivers/usb/usbip/vhci_hcd.c"],
            ["bug_006", "null pointer", "mm", "mm/filemap.c"],
            ["bug_007", "null pointer", "media", "drivers/media/common/videobuf2/videobuf2-core.c"],
            ["bug_008", "null pointer", "fs", "fs/namespace.c"],
        ],
        [1700, 1800, 1500, 4360],
        "表4-1 正式实验样本",
    )

    add_heading(doc, "4.3 轻量静态检查", 2)
    add_body(doc, "check_patch.py检查输出JSON是否有效、original是否能在片段中匹配、是否修改定位文件、是否出现大段删除、提前return、goto或注释核心代码，并检查reason和semantic preservation reason。该检查用于发现明显格式和风险问题，但不理解完整控制流，因此只能作为预筛选。")

    add_heading(doc, "4.4 分层评价指标", 2)
    add_body(doc, "本文将评价分为人工语义、真实源码适用、编译和动态验证四层。较低层覆盖更多输出但证据较弱；较高层证据更强但资源成本更高。未执行的实验记为not_run，历史未记录的成本与重试记为unknown，二者均不作为失败计数。")
    add_picture(doc, figures["evidence"], "图4-1 从人工评价到动态复现的证据链", width=6.1)
    add_table(
        doc,
        ["指标", "定义", "判定边界"],
        [
            ["人工语义标签", "对照开发者补丁进行plausible/helpful/incorrect判断", "不是动态修复率"],
            ["定位命中", "模型编辑文件是否覆盖开发者修复文件", "当前Prompt已披露文件，仅表示编辑目标命中"],
            ["补丁应用率", "模型edit能否在真实parent commit应用并形成diff", "通过不代表语义正确"],
            ["编译通过率", "应用补丁后指定内核目标能否构建", "局部编译不覆盖全部配置"],
            ["崩溃消除率", "同一reproducer下目标崩溃是否消失", "有限时间未复现不等于完全正确"],
            ["开发者补丁相似度", "模型与developer changed lines的词法相似度", "补充指标，不能替代语义评价"],
            ["API成本与重试", "依据usage和运行时价格记录调用资源", "历史响应未保存usage时记unknown"],
        ],
        [2200, 3800, 3360],
        "表4-2 扩展评价指标及边界",
    )

    add_heading(doc, "4.5 工程验证实现", 2)
    add_body(doc, "真实源码核验在每个developer patch的parentOfFixCommit上执行checkout、git apply --check、模型替换和diff检查。局部编译对bug_008的parent、developer patch和LLM improved三个版本使用同一工具链、defconfig和fs/namespace.o目标。动态验证使用kGymSuite、Docker、QEMU/KVM和syzkaller C reproducer，先运行parent确认目标崩溃可复现，再验证patched image。")

    add_heading(doc, "5 实验结果与分析", 1)
    add_heading(doc, "5.1 人工语义评价", 2)
    add_body(doc, "24个模型输出均完成plausible、helpful、incorrect三分类。整体得到plausible 2个、helpful 7个、incorrect 15个，plausible与helpful合计9/24。Improved组有用输出为4/8，高于Baseline的2/8和With Trace的3/8。")
    add_picture(doc, figures["groups"], "图5-1 三组人工语义评价分布", width=6.0)
    add_table(
        doc,
        ["组别", "plausible", "helpful", "incorrect", "plausible+helpful"],
        [
            ["Baseline", "1/8", "1/8", "6/8", "2/8（25.0%）"],
            ["With Trace", "0/8", "3/8", "5/8", "3/8（37.5%）"],
            ["Improved", "1/8", "3/8", "4/8", "4/8（50.0%）"],
        ],
        [1800, 1500, 1500, 1500, 3060],
        "表5-1 人工评价分组结果",
    )
    add_body(doc, "Trace信息主要把部分incorrect输出提升为helpful，但未增加plausible数量。这说明短轨迹能够帮助模型理解崩溃方向，却不足以恢复开发者补丁中的完整错误处理语义。Semantic Guard组的useful比例最高，但8个样本的规模不足以进行显著性推断，因此本文将结果解释为小样本趋势与案例证据。")

    add_heading(doc, "5.2 根因与缺陷类型", 2)
    add_body(doc, "越界访问类最容易获得helpful输出：9个输出中7个为helpful，模型通常能从数组索引、长度或移位操作识别边界风险。空指针类包含2个plausible输出，均来自bug_008的操作顺序调整。内存泄漏类最困难，2个样本、3组共6个输出全部incorrect，反映了资源所有权和cleanup路径对局部上下文的依赖。")
    add_table(
        doc,
        ["缺陷类型", "输出数", "plausible", "helpful", "incorrect"],
        [
            ["memory leak", "6", "0", "0", "6"],
            ["out-of-bounds", "9", "0", "7", "2"],
            ["null pointer", "9", "2", "0", "7"],
        ],
        [2600, 1300, 1700, 1700, 2060],
        "表5-2 不同缺陷类型的人工评价",
    )

    add_heading(doc, "5.3 真实源码适用性", 2)
    add_body(doc, "8个developer patch全部能在对应parent commit通过apply检查，证明样本和commit映射正确。24个模型补丁中10个能真实应用：Baseline 3/8、With Trace 2/8、Improved 5/8。Improved较With Trace提高37.5个百分点，主要原因是Semantic Guard促使模型生成更明确的最小局部修改，同时减少无编辑输出。")
    add_table(
        doc,
        ["组别", "apply通过", "apply失败", "通过率", "平均diff相似度"],
        [
            ["Baseline", "3", "5", "37.5%", "0.1848"],
            ["With Trace", "2", "6", "25.0%", "0.0998"],
            ["Improved", "5", "3", "62.5%", "0.4078"],
        ],
        [2100, 1500, 1500, 1800, 2460],
        "表5-3 扩展指标分组摘要",
    )
    add_body(doc, "定位命中在所有提出edit的输出上均为100%，但该数字不能解释为端到端fault localization准确率，因为Prompt已经向模型提供localization file。它只能说明模型没有偏离给定文件。未来应在不披露文件的条件下计算top-k文件或函数定位率。")

    add_heading(doc, "5.4 编译与动态验证", 2)
    add_body(doc, "编译和动态验证只覆盖代表案例bug_008 improved。parent、developer和LLM improved三个版本均完成defconfig和fs/namespace.o局部编译；developer与LLM版本的源码diff和目标对象SHA-256一致，说明模型生成了与开发者相同的重排。")
    add_table(
        doc,
        ["版本", "patch apply", "defconfig", "局部编译", "对象哈希前缀"],
        [
            ["parent", "不适用", "pass", "pass", "8abccf3ac28a"],
            ["developer", "pass", "pass", "pass", "bad2f34483c0"],
            ["LLM improved", "pass", "pass", "pass", "bad2f34483c0"],
        ],
        [1900, 1800, 1500, 1700, 2460],
        "表5-4 bug_008三路局部编译结果",
    )
    add_body(doc, "动态阶段中，parent版本约59秒复现“KASAN: null-ptr-deref Read in sys_mount_setattr”。修复版本复用相同内核配置与reproducer，连续运行3次，每次6分钟，均由syz-crush正常结束，未出现目标崩溃、特殊崩溃或worker异常。该结果为bug_008提供端到端证据，但不能外推到其余23个输出，也不能证明长期调度下绝对无误。")

    add_heading(doc, "5.5 API成本与重试记录", 2)
    add_body(doc, "历史DeepSeek响应没有保存usage字段，因此24个既有输出的API成本和重试次数记为unknown，而非0。项目已扩展run_metadata.json：后续运行将记录三阶段call count、prompt tokens、completion tokens、执行时间、输入输出单价、总成本和retry count。价格必须使用运行时供应商公开价格，避免用当前价格回填历史成本。")

    add_heading(doc, "6 代表案例分析", 1)
    add_heading(doc, "6.1 bug_003：方向正确但不等价的越界修复", 2)
    add_body(doc, "bug_003的developer patch将text[len]改为r->text_buf[len]，根因是写入了错误缓冲区。三组模型均倾向于增加len < buf_size条件。该修改能够避免部分越界写入，因此被评为helpful，但改变了满缓冲区时的终止行为，且没有修复错误目标缓冲区问题。该案例说明“能阻止sanitizer报警”与“恢复开发者语义”并不等价。")

    add_heading(doc, "6.2 bug_005：边界检查与错误路径语义", 2)
    add_body(doc, "bug_005涉及1 << wValue在wValue大于等于32时产生未定义移位。With Trace和Improved均加入wValue < 32保护，正确识别了边界风险；developer patch则在非法输入时goto error。模型补丁选择静默跳过操作，可能使调用方观察到不同错误行为，因此只能评为helpful。Semantic Guard保存了合法路径，却未保证非法路径与内核API约定一致。")

    add_heading(doc, "6.3 bug_008：与开发者补丁一致的操作重排", 2)
    add_body(doc, "bug_008的崩溃源于namespace_unlock之后再调用cleanup_group_ids，导致mnt可能失效。Baseline和Improved都把cleanup移动到unlock之前，与developer patch一致。Improved还解释了锁释放后对象生命周期变化，以及重排不删除任何cleanup功能。该补丁在真实源码上可应用、局部编译通过，并完成动态验证，是本文证据最完整的正面案例。")

    add_heading(doc, "6.4 内存泄漏失败：局部约束不能恢复所有权", 2)
    add_body(doc, "bug_001的developer patch在硬件offload不受支持时将hook_list恢复到原列表；模型却围绕无关kmalloc检查生成no-op。bug_002的developer patch在generic_rw失败时释放frontend资源；Improved只在attach后增加NULL检查。两个案例都需要理解资源在多个函数和错误路径之间的所有权变化。Semantic Guard可以约束补丁形式，却不能从缺失上下文中推导真实cleanup责任，因此6个输出全部incorrect。")

    add_heading(doc, "7 讨论", 1)
    add_heading(doc, "7.1 对研究问题的回答", 2)
    add_body(doc, "对于RQ1，Baseline仅有2/8个plausible或helpful输出，说明局部崩溃信息不足以稳定完成内核修复，但模型能够在部分边界类问题上提供诊断方向。对于RQ2，Trace使有用输出从2/8增至3/8，主要提升helpful而非plausible。对于RQ3，Improved达到4/8有用输出和5/8真实源码apply，通过率均为三组最高，表明语义约束具有积极趋势。对于RQ4，只有1个输出完成编译和动态验证，证明不同证据层必须分别报告。")

    add_heading(doc, "7.2 与RGym的关系及可继续完善方向", 2)
    add_body(doc, "RGym表明本地化、调用栈、blamed commit和反馈重试可以在普通硬件上形成更现实的内核APR评测。本文已吸收其指标思想，但尚未实现完整反馈循环。下一步可以在现有三组之外增加反馈重试组：首次apply或compile失败后，把错误日志反馈给模型，只允许一次或两次受控修改，并记录每次成本、tokens、错误类别和最终状态。")
    add_body(doc, "定位方面，应从当前“给定文件后编辑是否命中”升级为隐藏localization file的top-1/top-5文件与函数定位实验；上下文方面，可检索blamed commit、相似历史补丁和调用图；验证方面，应逐步扩大局部编译和reproducer覆盖，而不是仅增加人工标签数量。")

    add_heading(doc, "7.3 有效性威胁", 2)
    add_body(doc, "内部有效性方面，trace summary由人工整理，可能包含先验偏差；人工标签依赖对developer patch的解释，虽有固定标准，仍存在主观性。构念有效性方面，diff相似度只能度量词法接近，apply通过只能证明编辑可落地，均不能代替行为正确性。外部有效性方面，8个样本和单一模型不足以代表所有内核子系统或LLM。结论有效性方面，小样本不适合显著性检验，本文仅报告描述性趋势。")

    add_heading(doc, "7.4 工程与伦理边界", 2)
    add_body(doc, "自动生成的内核补丁可能引入安全退化、竞态或静默功能损失，不应在缺少审查和测试的情况下进入生产内核。模型输出应作为候选建议，由开发者结合代码审查、静态分析、编译矩阵、reproducer和回归测试验证。API密钥、崩溃日志和未公开漏洞信息也应遵守最小披露原则。")

    add_heading(doc, "8 结论与展望", 1)
    add_body(doc, "本文实现了一个面向真实Linux内核崩溃的轻量化LLM修复实验系统。系统使用根因假设、自反思和补丁生成三阶段流程，并通过Semantic Guard约束功能破坏型修改。8个kBenchSyz样本的24个输出显示，Trace和Semantic Guard能够改善部分输出质量，其中Improved组plausible+helpful为50.0%，真实源码补丁适用率为62.5%。bug_008形成了从人工判断、真实源码apply、局部编译到动态复现的完整证据链。")
    add_body(doc, "同时，实验清楚揭示了方法边界：总体incorrect仍占62.5%，内存泄漏样本全部失败，绝大多数输出尚未完成编译与动态验证。因而，本文结论不是“LLM已经能够自动修复Linux内核”，而是“结构化推理、语义约束和分层验证可以提高候选补丁的可用性与可审计性”。未来应结合RGym式定位与反馈、Code Researcher式历史上下文和Live-kBench式时序隔离，扩大模型、样本和动态验证覆盖，并以成本可控的方式构建更可信的内核修复代理。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Le Goues C, Nguyen T, Forrest S, et al. GenProg: A Generic Method for Automatic Software Repair. IEEE Transactions on Software Engineering, 2012.",
        "[2] Lutellier T, Pham H V, Pang L, et al. CoCoNuT: Combining Context-Aware Neural Translation Models Using Ensemble for Program Repair. ISSTA, 2020.",
        "[3] Sobania D, Briesch M, Hanna C, et al. An Analysis of the Automatic Bug Fixing Performance of ChatGPT. APR Workshop, 2023.",
        "[4] Jimenez C E, Yang J, Wettig A, et al. SWE-bench: Can Language Models Resolve Real-World GitHub Issues? ICLR, 2024.",
        "[5] Mathai A, Huang C, Maniatis P, et al. kGym: A Platform and Dataset to Benchmark Large Language Models on Linux Kernel Crash Resolution. NeurIPS Datasets and Benchmarks Track, 2024.",
        "[6] Mathai A, Huang C, Ma S, et al. CrashFixer: A Crash Resolution Agent for the Linux Kernel. arXiv:2504.20412, 2025.",
        "[7] Shehada K, Wu Y, Feng W D, et al. Rethinking Kernel Program Repair: Benchmarking and Enhancing LLMs with RGym. NeurIPS 2025 LLM Evaluation Workshop, 2025.",
        "[8] Singh R, Joel S, Mehrotra A, et al. Code Researcher: Deep Research Agent for Large Systems Code and Commit History. arXiv:2506.11060, 2025.",
        "[9] Pabba A, Mathai A, Chakraborty A, et al. SemAgent: A Semantics Aware Program Repair Agent. arXiv:2506.16650, 2025.",
        "[10] Huang C, Mathai A, Yu F, et al. Outrunning LLM Cutoffs: A Live Kernel Crash Resolution Benchmark for All. arXiv:2602.02690, 2026.",
        "[11] Google. syzkaller: An Unsupervised Coverage-Guided Kernel Fuzzer. https://github.com/google/syzkaller.",
        "[12] ARiSE-Lab. kGymSuite: Evaluating Linux Patches at Scale. https://github.com/ARiSE-Lab/kGymSuite.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Pt(-22)
        p.paragraph_format.left_indent = Pt(22)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(ref)
        set_run_font(run, size=9.8)

    doc.add_page_break()
    add_heading(doc, "附录A 复现实验要点", 1)
    add_body(doc, "完整复现说明见项目docs/reproduction_guide.md。最小流程包括：准备8个固定样本；运行三组模型实验；执行replace-based静态检查；填写evaluation_real.csv；在真实parent commit上核验补丁；生成扩展指标；对代表案例执行局部编译和动态验证。重复调用模型前应备份正式outputs，并使用--output-root隔离mock结果。")
    add_table(
        doc,
        ["产物", "位置", "用途"],
        [
            ["人工评价", "results/evaluation_real.csv", "24个输出语义标签"],
            ["源码核验", "results/kernel_verify_summary.csv", "真实parent commit apply结果"],
            ["扩展指标", "results/extended_metrics.csv", "定位、相似度、成本与重试"],
            ["编译对照", "results/local_compile_comparison/", "bug_008三路局部编译证据"],
            ["动态验证", "results/dynamic_validation/", "parent复现与patched重复验证"],
        ],
        [1900, 3900, 3560],
        "表A-1 主要实验产物",
    )


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "flow": ASSET_DIR / "paper_flow.png",
        "groups": ASSET_DIR / "group_results.png",
        "evidence": ASSET_DIR / "evidence_chain.png",
    }
    draw_flow(figures["flow"])
    draw_group_chart(figures["groups"])
    draw_evidence_chart(figures["evidence"])

    doc = setup_document()
    add_cover(doc)
    add_abstracts(doc)
    add_contents(doc)
    build_content(doc, figures)

    core = doc.core_properties
    core.title = "LLM辅助Linux内核缺陷修复"
    core.subject = "软件质量保证课程结项论文"
    core.author = ""
    core.keywords = "LLM, Linux kernel, automated program repair, Semantic Guard"
    core.comments = "Generated from the project's verified experiment artifacts."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
